"""Deterministic PDF/DOCX extraction with explicit failures and page provenance."""

from __future__ import annotations

import math
import re
import unicodedata
from collections import Counter
from dataclasses import dataclass
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader

from operational_decision.rag.document_catalog import DocumentDescriptor


class ExtractionError(RuntimeError):
    """Report a controlled EXTRACTION_FAILED document outcome."""

    code = "EXTRACTION_FAILED"


@dataclass(frozen=True)
class LoadedPage:
    """Extracted and cleaned text for one source page."""

    page_number: int
    content: str


@dataclass(frozen=True)
class LoadedDocument:
    """One manifest document with page-level extracted content."""

    descriptor: DocumentDescriptor
    pages: tuple[LoadedPage, ...]


def deterministic_cleanup(text: str, repeated_edge_lines: set[str] | None = None) -> str:
    """Normalize text without removing meaningful numbering or legal structure."""
    normalized = unicodedata.normalize("NFKC", text).replace("\u00ad", "")
    lines: list[str] = []
    blank = False
    for raw_line in normalized.splitlines():
        line = re.sub(r"[\t \f\v]+", " ", raw_line).strip()
        if repeated_edge_lines and line in repeated_edge_lines:
            continue
        if not line:
            if lines and not blank:
                lines.append("")
            blank = True
            continue
        lines.append(line)
        blank = False
    return "\n".join(lines).strip()


def _repeated_edge_lines(raw_pages: list[str]) -> set[str]:
    if len(raw_pages) < 2:
        return set()
    counts: Counter[str] = Counter()
    for text in raw_pages:
        lines = [re.sub(r"\s+", " ", line).strip() for line in text.splitlines()]
        nonempty = [line for line in lines if line]
        counts.update(set(nonempty[:3] + nonempty[-3:]))
    threshold = max(2, math.ceil(len(raw_pages) * 0.6))
    return {line for line, count in counts.items() if count >= threshold and len(line) < 180}


class DocumentLoader:
    """Extract supported local documents without OCR or network fallback."""

    def load(self, descriptor: DocumentDescriptor, path: Path) -> LoadedDocument:
        """Load a PDF or DOCX and fail openly when usable text is unavailable."""
        suffix = path.suffix.casefold()
        if suffix == ".pdf":
            pages = self._load_pdf(path)
        elif suffix == ".docx":
            pages = self._load_docx(path)
        else:
            raise ExtractionError(f"EXTRACTION_FAILED: unsupported type: {path.suffix}")
        if not pages or not any(page.content.strip() for page in pages):
            raise ExtractionError(f"EXTRACTION_FAILED: no text extracted from {path.name}")
        return LoadedDocument(descriptor=descriptor, pages=tuple(pages))

    @staticmethod
    def _load_pdf(path: Path) -> list[LoadedPage]:
        try:
            reader = PdfReader(path)
            if reader.is_encrypted:
                raise ExtractionError(f"EXTRACTION_FAILED: encrypted PDF: {path.name}")
            raw_pages = [page.extract_text() or "" for page in reader.pages]
        except ExtractionError:
            raise
        except Exception as error:
            raise ExtractionError(f"EXTRACTION_FAILED: cannot read PDF: {path.name}") from error
        repeated = _repeated_edge_lines(raw_pages)
        return [
            LoadedPage(index, deterministic_cleanup(text, repeated))
            for index, text in enumerate(raw_pages, start=1)
        ]

    @staticmethod
    def _load_docx(path: Path) -> list[LoadedPage]:
        try:
            document = DocxDocument(str(path))
            blocks = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    blocks.append(" | ".join(cell.text for cell in row.cells))
        except Exception as error:
            raise ExtractionError(f"EXTRACTION_FAILED: cannot read DOCX: {path.name}") from error
        return [LoadedPage(1, deterministic_cleanup("\n\n".join(blocks)))]